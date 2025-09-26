import pandas as pd
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import os
import random

# --- 1. 文件路径 (请根据您的实际情况修改) ---
# CSV文件路径，包含误差数据
csv_path = r"D:\code\polyline\output\per_row_errors.csv"
# 图片文件夹路径
image_folder = r"D:\code\polyline\output\ditch_origin"
# 生成的Word报告的保存路径
output_word_path = r"D:\code\polyline\output\ditch_report_with_summary_v3.docx"

# --- 2. 读取并准备CSV数据 ---
try:
    print(f"正在读取CSV文件: {csv_path}")
    df = pd.read_csv(csv_path)
    sort_column = '绝对百分比误差(%)'
    if sort_column not in df.columns:
        raise ValueError(f"CSV文件缺少排序列 '{sort_column}'")

    # 额外的数据清洗：确保误差列是数值类型
    df[sort_column] = pd.to_numeric(df[sort_column], errors='coerce')
    df.dropna(subset=[sort_column], inplace=True)

    # 为后续的详细列表对数据框进行排序
    df_sorted = df.sort_values(by=sort_column, ascending=False)
    print("CSV读取并排序成功。")
except FileNotFoundError:
    print(f"错误: CSV文件未找到于路径 '{csv_path}'。请检查文件路径是否正确。")
    exit()
except Exception as e:
    print(f"读取或排序CSV时发生错误: {e}")
    exit()

# --- 3. 创建Word文档 ---
doc = docx.Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'  # 建议使用支持中文的字体，如'等线'或'微软雅黑'
font.size = Pt(10.5)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(8)

# --- 4. 在文档开头添加摘要部分 ---
doc.add_heading('清沟长度对比分析报告', level=0)
run = doc.add_paragraph().add_run(f"报告生成时间: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
run.font.size = Pt(9)
run.italic = True
#
# doc.add_heading('误差区间分布统计', level=1)
#
# total_count = len(df)
# if total_count > 0:
#     # --- 定义用于分组的误差区间和标签 ---
#     bins = [0, 10, 20, 30, 40, 50, 60, 70, 80, 90, 100, float('inf')]
#     # <-- 问题修正：移除了重复的 '(90, 100]' 标签
#     labels = [f"({bins[i]}, {bins[i + 1]}]" for i in range(len(bins) - 2)] + ["> 100"]
#
#     df['error_bin'] = pd.cut(df[sort_column], bins=bins, labels=labels, right=True)
#
#     # --- 创建包含新列的摘要表格 ---
#     summary_table = doc.add_table(rows=1, cols=4)
#     summary_table.style = 'Table Grid'
#     hdr_cells = summary_table.rows[0].cells
#     hdr_cells[0].text = '误差范围 (%)'
#     hdr_cells[1].text = '数量'
#     hdr_cells[2].text = '区间占比 (%)'
#     hdr_cells[3].text = '示例图片'
#
#     # --- 按误差区间分组并填充表格 ---
#     # 使用 observed=True 来确保只处理数据中实际存在的区间
#     grouped = df.groupby('error_bin', observed=True)
#
#     for bin_label, group_df in grouped:
#         count = len(group_df)
#         if count == 0:
#             continue  # 跳过没有数据的区间
#
#         percentage = (count / total_count) * 100
#
#         row_cells = summary_table.add_row().cells
#         row_cells[0].text = str(bin_label)
#         row_cells[1].text = str(count)
#         row_cells[2].text = f"{percentage:.2f} %"
#
#         # --- 在最后一列嵌入示例图片 ---
#         image_cell = row_cells[3]
#         # 设置单元格内容垂直居中
#         image_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#
#         # 从分组中抽取最多2个样本
#         sample_images = group_df.head(2)
#
#         # 使用单元格中已存在的第一个段落
#         p = image_cell.paragraphs[0]
#         p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#
#         image_added = False
#         for _, row in sample_images.iterrows():
#             # 需要在段落内创建一个新的 Run 对象来添加图片
#             run = p.add_run()
#             image_name = f"ditch__{row['name']}__{row['CODE']}__proj.png"
#
#             image_path = os.path.join(image_folder, image_name)
#
#             if os.path.exists(image_path):
#                 if image_added:
#                     # 在图片之间添加空格
#                     p.add_run('  ')
#                 # 添加图片，并调整大小以适应单元格
#                 run.add_picture(image_path, width=Inches(4.0))
#                 image_added = True
#
#         if not image_added:
#             p.add_run("无可用图片")
#
#     print("已生成误差区间分布统计表。")
# else:
#     doc.add_paragraph("无有效数据可用于生成误差统计。")

# 添加分页符，将摘要与详细数据分开
doc.add_page_break()

# --- 5. 遍历数据，生成详细报告 ---
doc.add_heading('清沟误差详细列表', level=1)
for index, row in df_sorted.iterrows():
    try:
        ditch_name = row['name']
        ditch_code = row['CODE']

        print(f"正在处理: {ditch_name}, CODE: {ditch_code}")
        doc.add_heading(f"清沟: {ditch_name} (CODE: {ditch_code})", level=2)

        # 建表（展示关键数据）
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        keys_to_show = {
            '清沟实际长度': f"{row.get('清沟实际长度', 0):.2f}",
            '堤坝线投影长度': f"{row.get('堤坝线投影长度', 0):.2f}",
            '人工投影长度': f"{row.get('人工投影长度', 0):.2f}",
            '绝对百分比误差(%)': f"{row.get(sort_column, 0):.2f} %"
        }
        for i, (key, value) in enumerate(keys_to_show.items()):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = value

        # 插图逻辑
        image_name = f"ditch__{ditch_name}__{ditch_code}__proj.png"
        image_with_closedshape_name=f"ditch__{ditch_name}__{ditch_code}__closed.png"
        image_path = os.path.join(image_folder, image_name)
        image_with_closedshape_path=os.path.join(image_folder,image_with_closedshape_name)

        if os.path.exists(image_path):
            # 添加一个空段落以增加图片和表格之间的间距
            doc.add_paragraph()
            # 插入图片并居中
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(image_path, width=Inches(5.5))
            run.add_picture(image_with_closedshape_path, width=Inches(5.5))
            print(f"  - 图片 '{image_name}' 已添加。")
        else:
            p = doc.add_paragraph(f"警告: 未找到图片文件 '{image_name}'")
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # 红色字体
            print(f"  - 警告: 未找到图片 '{image_path}'")

    except KeyError as e:
        print(f"处理行 {index} 时发生键错误: 缺少列 {e}。已跳过此行。")
        continue
    except Exception as e:
        print(f"处理行 {index} 时发生未知错误: {e}。已跳过此行。")
        continue

# --- 6. 保存Word ---
try:
    doc.save(output_word_path)
    print("\n-----------------------------------------")
    print(f"处理完成！报告已保存到: {output_word_path}")
    print("-----------------------------------------")
except Exception as e:
    print(f"保存Word文档时发生错误: {e}")
