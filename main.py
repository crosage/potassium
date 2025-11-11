import os
import geopandas as gpd
import pandas as pd
import numpy as np
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from processing.close_shape import generate_closed_shapes_with_polylines
from processing.normals import generate_infinite_normals_on_linestring_with_polyline
from file_io.file_io import (
    load_polylines_from_shp, load_left_right_lines_from_json, save_left_right_lines_to_json,
    load_merged_polyline_from_json, save_merged_polyline_to_json, load_split_points_from_file,
    save_split_points_to_file, load_closed_shapes_from_file, save_closed_shapes_to_file,
)
from processing.ditch import process_ditch_endpoints
from processing.merging import merge_polylines, plot_merge_verification
from processing.splitting import split_polyline_by_points, extract_subcurve, preprocess_crop_lines
from shapely.geometry import LineString, Point

from visualization.plot import plot_closed_shapes, plot_normals, plot_river_elements, plot_river_with_satellite


# ==============================================================================
# SECTION 1: 核心地理空间处理函数
# ==============================================================================

def check_file_exists(file_path):
    """检查文件是否存在"""
    return os.path.exists(file_path)


def orient_line_by_coordinate(line: LineString) -> LineString:
    """
    根据坐标矫正线的方向，确保其起点总是X坐标较小（或Y坐标较小）的一端。
    """
    if not isinstance(line, LineString) or len(line.coords) < 2:
        return line

    start_point = line.coords[0]
    end_point = line.coords[-1]

    if start_point[0] > end_point[0]:
        return LineString(list(line.coords)[::-1])
    elif abs(start_point[0] - end_point[0]) < 1e-6 and start_point[1] > end_point[1]:
        return LineString(list(line.coords)[::-1])
    else:
        return line


def run_processing_pipeline(config):
    """
    运行完整的地理数据处理流水线。
    现在使用北侧大堤线和南侧大堤线进行河道划分。
    """
    job_name = config["job_name"]
    output_dir = config["output_dir"]
    plot_cfg = config.get("plot_config")

    print(f"\n--- 开始处理任务: {job_name} ---")
    os.makedirs(output_dir, exist_ok=True)

    # 路径定义
    merged_polyline_path = os.path.join(output_dir, f"merged_polyline.json")
    normals_path = os.path.join(output_dir, f"normals.json")
    closed_shapes_path = os.path.join(output_dir, f"closed_shapes.json")
    ditch_output_path_prefix = os.path.join(output_dir, f"ditch_{job_name}")

    # 1. 加载中心线
    print("1. 加载中心线数据...")
    centerlines = load_polylines_from_shp(config["centerline_path"])

    # 2. 合并中心线
    if check_file_exists(merged_polyline_path):
        print(f"2. 合并中心线... 检测到缓存文件，跳过。({merged_polyline_path})")
        merged_polyline = load_merged_polyline_from_json(merged_polyline_path)
    else:
        print("2. 合并中心线... 开始计算。")
        merged_polyline = merge_polylines(centerlines, log=False)
        save_merged_polyline_to_json(merged_polyline, merged_polyline_path)
    merged_polyline = merged_polyline.line

    # 3. 加载南北大堤线（替代原来的左右岸线）
    print("3. 加载南北大堤线数据...")
    north_dam_polylines = load_polylines_from_shp(config["north_dam_line"], ignore=False, charset="GBK")
    south_dam_polylines = load_polylines_from_shp(config["south_dam_line"], ignore=False, charset="GBK")

    # 合并并矫正方向
    north_dam_line = merge_polylines(north_dam_polylines, False).line
    south_dam_line = merge_polylines(south_dam_polylines, False).line
    north_dam_line = orient_line_by_coordinate(north_dam_line)
    south_dam_line = orient_line_by_coordinate(south_dam_line)
    print("南北大堤线加载完成。")

    # 4. 裁剪线（如果需要的话，使用南北大堤线）
    print("4. 裁剪中心线和大堤线，防止出现错误相连")
    # 如果 preprocess_crop_lines 函数需要调整，可能需要修改参数
    centerline, north_dam_line, south_dam_line = preprocess_crop_lines(
        merged_polyline, north_dam_line, south_dam_line
    )

    # 5. 生成法线（使用北侧和南侧大堤线）
    if check_file_exists(normals_path):
        print(f"5. 生成法线... 检测到缓存文件，跳过。({normals_path})")
        normals = load_split_points_from_file(normals_path)
    else:
        print("5. 生成法线... 开始计算。")
        normals = generate_infinite_normals_on_linestring_with_polyline(
            centerline, north_dam_line, south_dam_line, interval=config["normal_interval"]
        )
        save_split_points_to_file(normals, normals_path)

    # 6. 封闭形状生成（使用北侧和南侧大堤线）
    if check_file_exists(closed_shapes_path):
        print(f"6. 封闭形状生成... 检测到缓存文件，跳过。({closed_shapes_path})")
        closed_shapes = load_closed_shapes_from_file(closed_shapes_path)
    else:
        print("6. 封闭形状生成... 开始计算。")
        closed_shapes = generate_closed_shapes_with_polylines(
            normals, north_dam_line, south_dam_line, meters=config["shape_max_length"], log=False
        )
        save_closed_shapes_to_file(closed_shapes, closed_shapes_path)

    # 7. 处理沟渠数据
    print("7. 处理沟渠数据...")
    ditches = load_polylines_from_shp(config["ditch_path"], ignore=False)
    manual_shp_path = config.get("manual_ditch_path")

    # # 绘制卫星影像叠加图（使用南北大堤线）
    # plot_river_with_satellite(
    #     left_line=north_dam_line,  # 南侧大堤线（作为参考）
    #     right_line=south_dam_line,  # 南侧大堤线
    #     ditches=ditches,
    #     center_line=centerline,
    #     title="黄河研究区域卫星影像图 (2025-01-22)",
    #     save_path='output/overview_20250122_full.png',
    #     save_clean_path='output/overview_20250122_clean.png',
    #     show_axis=True,
    #     use_proxy=True
    # )

    # 调用处理函数（使用南北大堤线）
    process_ditch_endpoints(
        ditchs=ditches,
        closed_shapes=closed_shapes,
        north_dam_line=north_dam_line,  # 传入北侧大堤线
        south_dam_line=south_dam_line,  # 传入南侧大堤线
        centerline=merged_polyline,
        save_path=ditch_output_path_prefix,
        log=True,
        manual_shp_path=manual_shp_path
    )
    print(f"--- 任务: {job_name} 处理完成 ---")

# ==============================================================================
# SECTION 2: 报告与分析函数
# ==============================================================================

def generate_error_reports(auto_csv_path: str, out_dir: str):
    """
    根据沟渠处理结果计算误差指标，并保存详细和汇总的CSV文件。
    """
    print("\n--- 开始生成误差报告 ---")
    if not os.path.exists(auto_csv_path):
        print(f"错误: 在'{auto_csv_path}'未找到输入CSV文件。跳过报告生成。")
        return

    os.makedirs(out_dir, exist_ok=True)
    df = pd.read_csv(auto_csv_path)

    # 计算逐行误差
    df["error_m"] = df["堤坝线投影长度"] - df["人工投影长度"]
    df["abs_error_m"] = df["error_m"].abs()
    df["squared_error_m2"] = df["error_m"] ** 2
    den = np.maximum(df["人工投影长度"], df["堤坝线投影长度"])
    df["绝对百分比误差(%)"] = (df["abs_error_m"] / den) * 100.0
    per_row = df.copy()

    # 按 (CODE, RIVERPART) 聚合
    def agg_block(g):
        # **MODIFIED: Removed dependency on "name" column for aggregation.**
        return pd.Series({
            "records": len(g),
            "latest_date": pd.to_datetime(g["DATE"], errors="coerce").max(),
            "人工投影长度_mean": float(g["人工投影长度"].mean()),
            "堤坝线投影长度_mean": float(g["堤坝线投影长度"].mean()),
            "MAE_m": float(g["abs_error_m"].mean()),
            "RMSE_m": float(np.sqrt(g["squared_error_m2"].mean())),
            "MAPE_percent_macro_per_ditch": float(g["绝对百分比误差(%)"].mean(skipna=True)),
        })

    per_ditch = (per_row.groupby(["CODE", "RIVERPART"], as_index=False)
                 .apply(agg_block).reset_index(drop=True))

    # 计算全局指标
    global_mse = float(per_row["squared_error_m2"].mean())
    global_rmse = float(np.sqrt(global_mse))
    macro_mape_all = float(per_row["绝对百分比误差(%)"].mean(skipna=True))
    micro_den_series = np.maximum(per_row["人工投影长度"], per_row["堤坝线投影长度"])
    mask = (per_row["人工投影长度"] > 0) & per_row["堤坝线投影长度"].notna()
    micro_num = per_row.loc[mask, "abs_error_m"].sum()
    micro_den = micro_den_series[mask].sum()
    micro_mape_all = (micro_num / micro_den) * 100.0 if micro_den > 0 else np.nan

    # 导出到CSV
    per_row_path = os.path.join(out_dir, "per_row_errors.csv")
    per_ditch_path = os.path.join(out_dir, "per_ditch_summary.csv")
    per_row.to_csv(per_row_path, index=False, encoding='utf-8-sig')
    per_ditch.to_csv(per_ditch_path, index=False, encoding='utf-8-sig')

    print(f"✔ 详细误差报告已保存至: {per_row_path}")
    print(f"✔ 汇总报告已保存至: {per_ditch_path}")
    print(
        f"全局 (算法 vs 人工) RMSE: {global_rmse:.3f} m, Macro-MAPE: {macro_mape_all:.3f}%, Micro-MAPE: {micro_mape_all:.3f}%")
    print("--- 误差报告生成完成 ---")


def generate_word_report(csv_path: str, image_folder: str, output_word_path: str):
    """
    生成包含详细误差分析和图片的Word文档报告。
    新增功能：在报告顶部添加总体长度误差摘要。
    """
    print("\n--- 开始生成Word报告 ---")
    try:
        df = pd.read_csv(csv_path)
        sort_column = '绝对百分比误差(%)'
        if sort_column not in df.columns:
            raise ValueError(f"CSV文件缺少必需的排序列 '{sort_column}'")
        df[sort_column] = pd.to_numeric(df[sort_column], errors='coerce')
        df.dropna(subset=[sort_column], inplace=True)
        df_sorted = df.sort_values(by=sort_column, ascending=False)
        print(f"成功从以下路径加载并排序数据: {csv_path}")
    except FileNotFoundError:
        print(f"错误: 在'{csv_path}'未找到CSV文件。无法生成Word报告。")
        return
    except Exception as e:
        print(f"读取CSV时发生错误: {e}")
        return

    # --- 新增操作：计算总长度和总体误差 ---
    total_manual_length = df['人工投影长度'].sum()
    total_algo_length = df['堤坝线投影长度'].sum()
    total_error_m = total_algo_length - total_manual_length

    # 计算总体百分比误差，并处理分母为零的情况
    if total_manual_length > 0:
        total_percentage_error = (abs(total_error_m) / total_manual_length) * 100
    else:
        total_percentage_error = float('nan')

    # --- 创建Word文档 ---
    doc = docx.Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '等线'
    font.size = Pt(10.5)

    # --- 添加文档标题和时间戳 ---
    doc.add_heading('清沟长度对比分析报告', level=0)
    run = doc.add_paragraph().add_run(f"报告生成时间: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(9)
    run.italic = True

    # --- 新增操作：在Word中添加总体摘要表格 ---
    doc.add_heading('总体长度对比摘要', level=1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Table Grid'

    summary_data = {
        '人工清沟投影总长度': f"{total_manual_length:.2f} m",
        '算法计算总长度': f"{total_algo_length:.2f} m",
        '总长度误差': f"{total_error_m:.2f} m",
        '总体百分比误差(%)': f"{total_percentage_error:.2f} %" if not np.isnan(total_percentage_error) else "N/A"
    }

    for i, (key, value) in enumerate(summary_data.items()):
        summary_table.rows[i].cells[0].text = key
        summary_table.rows[i].cells[1].text = value

    # 使用分页符将摘要与详情分开
    doc.add_page_break()

    # --- 添加详细列表（此部分逻辑不变） ---
    doc.add_heading('清沟误差详细列表', level=1)
    for index, row in df_sorted.iterrows():
        try:
            # **MODIFIED: Use CODE and RIVERPART as primary identifiers, removing dependency on 'name'.**
            ditch_code = row['CODE']
            river_part = row['RIVERPART']
            # Create a unique, file-safe identifier consistent with the image generation script.
            unique_file_identifier = f"R_{ditch_code}_C_{river_part}"

            doc.add_heading(f"清沟 (CODE: {ditch_code}, RIVERPART: {river_part})", level=2)

            # 数据表格
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'
            keys_to_show = {
                '清沟实际长度': f"{row.get('清沟实际长度', 0):.2f} m",
                '堤坝线投影长度': f"{row.get('堤坝线投影长度', 0):.2f} m",
                '人工投影长度': f"{row.get('人工投影长度', 0):.2f} m",
                '绝对百分比误差(%)': f"{row.get(sort_column, 0):.2f} %"
            }
            for i, (key, value) in enumerate(keys_to_show.items()):
                table.rows[i].cells[0].text = key
                table.rows[i].cells[1].text = value

            # **MODIFIED: Construct image names using the new unique identifier.**
            image_name = f"ditch__{unique_file_identifier}__proj.png"
            image_with_closedshape_name = f"ditch__{unique_file_identifier}__closed.png"
            image_path = os.path.join(image_folder, image_name)
            image_with_closedshape_path = os.path.join(image_folder, image_with_closedshape_name)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if os.path.exists(image_path):
                p.add_run().add_picture(image_path, width=Inches(5.5))
            else:
                p.add_run(f"\n[警告: 图片未找到: {image_name}]").font.color.rgb = RGBColor(255, 0, 0)

            if os.path.exists(image_with_closedshape_path):
                p.add_run().add_picture(image_with_closedshape_path, width=Inches(5.5))
            else:
                p.add_run(f"\n[警告: 图片未找到: {image_with_closedshape_name}]").font.color.rgb = RGBColor(255, 0, 0)

        except KeyError as e:
            print(f"因缺少列 {e}，已跳过行 {index}。")
            continue
        except Exception as e:
            print(f"处理行 {index} 时发生未知错误: {e}。")
            continue

    # --- 保存Word文档 ---
    try:
        doc.save(output_word_path)
        print("\n-----------------------------------------")
        print(f"✔ Word报告成功保存至: {output_word_path}")
        print("-----------------------------------------")
    except Exception as e:
        print(f"保存Word文档时出错: {e}")

# ==============================================================================
# SECTION 3: 主执行模块
# ==============================================================================

def main():
    # --- 任务定义 ---
    tasks = [
        {
            "job_name": "20231226",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2024/20231226清沟_hz.shp"),
            "manual_ditch_path": os.path.join("data", "2024/20231226清沟映射长度_hz.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20240101",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2024/20240101清沟hz.shp"),
            "manual_ditch_path": os.path.join("data", "2024/20240101清沟映射长度hz.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20240107",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2024/20240107清沟_hz.shp"),
            "manual_ditch_path": os.path.join("data", "2024/20240107清沟映射长度_hz.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20240115",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2024/20240115清沟_hz.shp"),
            "manual_ditch_path": os.path.join("data", "2024/20240115清沟映射长度_hz.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20240123",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2024/20240123清沟_hz.shp"),
            "manual_ditch_path": os.path.join("data", "2024/20240123清沟映射长度_hz.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20240130",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2024/20240130清沟_hz.shp"),
            "manual_ditch_path": os.path.join("data", "2024/20240130清沟映射长度_hz.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20240205",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2024/20240205清沟_hz.shp"),
            "manual_ditch_path": os.path.join("data", "2024/20240205清沟映射长度_hz.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20240217",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2024/20240217清沟_hz.shp"),
            "manual_ditch_path": os.path.join("data", "2024/20240217清沟映射长度_hz.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20240225",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2024/20240225清沟_hz.shp"),
            "manual_ditch_path": os.path.join("data", "2024/20240225清沟映射长度_hz.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20240303",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2024/20240303清沟_hz.shp"),
            "manual_ditch_path": os.path.join("data", "2024/20240303清沟映射长度_hz.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20240307",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2024/20240307清沟_hz.shp"),
            "manual_ditch_path": os.path.join("data", "2024/20240307清沟映射长度_hz.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20240309",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2024/20240309清沟_hz.shp"),
            "manual_ditch_path": os.path.join("data", "2024/20240309清沟映射长度_hz.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20240311",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2024/20240311清沟_hz.shp"),
            "manual_ditch_path": os.path.join("data", "2024/20240311清沟映射长度_hz.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20240313",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2024/20240313清沟_hz.shp"),
            "manual_ditch_path": os.path.join("data", "2024/20240313清沟映射长度_hz.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20240317",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2024/20240317清沟_hz.shp"),
            "manual_ditch_path": os.path.join("data", "2024/20240317清沟映射长度_hz.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20250106",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2025/20250106/20250106清沟.shp"),
            "manual_ditch_path": os.path.join("data", "2025/20250106/20250106清沟映射长度.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20250115",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2025/20250115/清沟汇总_2024-2025年度20250115.shp"),
            "manual_ditch_path": os.path.join("data", "2025/20250115/c20240115专题监测清沟映射长度_Merge.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20250122",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2025/20250122/清沟汇总_2024-2025年度20250122.shp"),
            "manual_ditch_path": os.path.join("data", "2025/20250122/清沟映射汇总_2024-2025年度20250122.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
            "plot_config": {
                "normals": {
                    "x_min": 300000, "x_max": 320000,
                    "y_min": 4490000, "y_max": 4510000,
                    "rtree_mode": "point"   # 可选: "point" 或 "line"
                },
                "closed_shapes": {
                    "x_min": 300000, "x_max": 320000,
                    "y_min": 4490000, "y_max": 4510000
                }
            }
        },
        {
            "job_name": "20250203",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2025/20250203/清沟汇总_2024-2025年度20250203.shp"),
            "manual_ditch_path": os.path.join("data", "2025/20250203/清沟映射汇总_2024-2025年度20250203.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20250210",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2025/20250210/清沟汇总_2024-2025年度20250210.shp"),
            "manual_ditch_path": os.path.join("data", "2025/20250210/清沟映射汇总_2024-2025年度20250210.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20250213",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2025/20250213/FL_20250213清沟汇总.shp"),
            "manual_ditch_path": os.path.join("data", "2025/20250213/FL20250213fl清沟映射.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20250224",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2025/20250224/20250224_清沟_汇总.shp"),
            "manual_ditch_path": os.path.join("data", "2025/20250224/20250224_清沟映射_汇总.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20250302",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2025/20250302/清沟汇总_2024-2025年度20250302.shp"),
            "manual_ditch_path": os.path.join("data", "2025/20250302/清沟映射汇总_2024-2025年度20250302.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20250308",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2025/20250308/清沟汇总_2024-2025年度20250308.shp"),
            "manual_ditch_path": os.path.join("data", "2025/20250308/清沟映射汇总_2024-2025年度20250308.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20250309",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2025/20250309/清沟汇总_2024-2025年度202503091.shp"),
            "manual_ditch_path": os.path.join("data", "2025/20250309/清沟映射汇总_2024-2025年度202503091.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20250312",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2025/20250312/清沟汇总2024-2025_2025031.shp"),
            "manual_ditch_path": os.path.join("data", "2025/20250312/清沟映射汇总2024-2025_2025031.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20250315",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2025/20250315/清沟汇总2024-2025_20250315.shp"),
            "manual_ditch_path": os.path.join("data", "2025/20250315/清沟映射汇总2024-2025_20250315.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
        {
            "job_name": "20250317",
            "centerline_path": os.path.join("data","中心线平滑.shp"),
            "ditch_path": os.path.join("data", "2025/20250317/清沟汇总2024-2025_20250317.shp"),
            "manual_ditch_path": os.path.join("data", "2025/20250317/清沟映射汇总2024-2025_20250317.shp"),
            "output_dir": "output",
            "normal_interval": 100,
            "shape_max_length": 500,
            "north_dam_line": r"D:\code\polyline\data\2021内蒙古河段每日封河长度矢量-修订最新版20220124.shp",
            "south_dam_line": r"D:\code\polyline\data\2021大堤南岸完善.shp",
        },
    ]

    # --- 循环并为每个任务执行所有步骤 ---
    for task_config in tasks:
        # 步骤 1: 运行核心地理空间处理
        run_processing_pipeline(task_config)

        # 根据任务输出定义路径
        job_name = task_config["job_name"]
        output_dir = task_config["output_dir"]
        ditch_output_dir = os.path.join(output_dir, f"ditch_{job_name}")

        # 报告生成所需的路径
        initial_results_csv = os.path.join(ditch_output_dir, "ditch_results.csv")
        detailed_errors_csv = os.path.join(ditch_output_dir, "per_row_errors.csv")
        word_report_path = os.path.join(output_dir, f"分析报告_{job_name}.docx")

        # 步骤 2: 生成CSV误差报告
        generate_error_reports(
            auto_csv_path=initial_results_csv,
            out_dir=ditch_output_dir
        )

        # 步骤 3: 生成最终的Word文档报告
        generate_word_report(
            csv_path=detailed_errors_csv,
            image_folder=ditch_output_dir,  # 图片位于同一文件夹
            output_word_path=word_report_path
        )

    print("\n所有任务已完成。")


if __name__ == "__main__":
    main()